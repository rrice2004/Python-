# SEC_REVIEW: Command-line tool for performing software security reviews.
# Capture, update, and document software risk, policies, and integrations.
# Generate PDF or Word reports for record-keeping or manual review.

import json
import os
import sys
from datetime import datetime
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter
from reportlab.lib.enums import TA_CENTER
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

# ------------------- CONFIG -------------------

REVIEW_DIR = "reviews"
os.makedirs(REVIEW_DIR, exist_ok=True)

YES_NO_QUESTIONS = [
    "privacy_policy", "tos_reviewed", "data_sharing", "pentesting"
]

# Updated QUESTIONS with software_name added, and software_vendor renamed
QUESTIONS = [
    ("software_name", "Software:"),
    ("software_vendor", "Vendor:"),
    ("use_case", "Enter the requested use case:"),
    ("submitted_by", "Who submitted the request?"),
    ("date_submitted", "Date submitted (leave blank for today):"),
    ("reviewer", "Reviewer name:"),
    ("review_date", "Review date (leave blank for today):"),
    ("version", "Version:"),
    ("overview", "Overview / Purpose (2–3 sentences):"),
    ("privacy_policy", "Privacy policy reviewed? (Y/N)"),
    ("privacy_notes", "Key privacy notes:"),
    ("tos_reviewed", "TOS reviewed? (Y/N)"),
    ("tos_notes", "Key TOS notes:"),
    ("data_sharing", "Data shared with third parties? (Y/N)"),
    ("data_sharing_notes", "Summary of data sharing:"),
    ("data_collected", "Data types collected:"),
    ("data_storage", "Storage location & retention:"),
    ("pentesting", "Pentesting performed? (Y/N)"),
    ("pentest_details", "Frequency & provider:"),
    ("certifications", "Security certifications (SOC2, ISO27001, etc.):"),
    ("auth_controls", "Auth controls (SSO/MFA/etc.):"),
    ("encryption", "Encryption details:"),
    ("def_general", "Defender General Score:"),
    ("def_security", "Defender Security Score:"),
    ("def_compliance", "Defender Compliance Score:"),
    ("def_notes", "Defender score concerns:"),
    ("integrations", "API / software integrations:"),
    ("integration_data", "Data exchanged:"),
    ("integration_security", "Integration security considerations:"),
    ("risk_level", "Risk level (Low/Moderate/High):"),
    ("key_risks", "Key risks (bulleted):"),
    ("mitigations", "Mitigations:"),
    ("approval", "Approval status (Approved, Conditional, Denied):"),
    ("approver", "Approver name:"),
    ("approval_date", "Approval date:"),
]

# ------------------- HELPER FUNCTIONS -------------------

def safe_input(prompt):
    """Input wrapper to handle KeyboardInterrupt with bold red message."""
    try:
        return input(prompt).strip()
    except KeyboardInterrupt:
        print("")
        print("***Input cancelled by user. Returning to menu***")
        return None

def list_reviews():
    """List existing JSON reviews and allow selection."""
    files = [f for f in os.listdir(REVIEW_DIR) if f.endswith(".json")]
    if not files:
        print("")
        print("***No saved reviews found***")
        return None

    print("Available Reviews:")
    for i, f in enumerate(files, 1):
        print(f"{i}. {f.replace('.json', '')}")

    choice = safe_input("Select a review number or press Enter to cancel: ")
    if choice is None:
        return None
    if choice.isdigit() and 1 <= int(choice) <= len(files):
        return os.path.join(REVIEW_DIR, files[int(choice) - 1])
    return None

def load_review(path=None):
    if path:
        with open(path, "r") as f:
            return json.load(f), path
    return {}, None

def save_review(data, path):
    with open(path, "w") as f:
        json.dump(data, f, indent=2)

# ------------------- REVIEW FUNCTIONS -------------------

def start_review():
    """Start a new review from scratch."""
    data = {}
    for key, question in QUESTIONS:
        answer = safe_input(question + " ")
        if answer is None:
            return  # cancelled
        if answer:
            data[key] = answer
        else:
            data[key] = ""
        if key in ["date_submitted", "review_date"] and not data[key]:
            data[key] = str(datetime.now().date())

    filename = f"{data.get('software_name', 'review').replace(' ', '_')}.json"
    path = os.path.join(REVIEW_DIR, filename)
    save_review(data, path)
    print(f"\nNew review saved as {filename}.\n")

def continue_review():
    """Load and edit an existing review."""
    existing = list_reviews()
    if not existing:
        return
    data, path = load_review(existing)
    for key, question in QUESTIONS:
        current = data.get(key, "")
        prompt_text = f"{question} (current: {current}) " if current else question + " "
        answer = safe_input(prompt_text)
        if answer is None:
            return
        if answer:
            data[key] = answer
        elif not current:
            data[key] = ""
        if key in ["date_submitted", "review_date"] and not data[key]:
            data[key] = str(datetime.now().date())
    save_review(data, path)
    print(f"\nReview updated: {os.path.basename(path)}\n")

# ------------------- PDF GENERATION -------------------

def generate_pdf():
    """Generate PDF with wrapped table cells."""
    existing = list_reviews()
    if not existing:
        return
    data, _ = load_review(existing)

    software = data.get("software_name", "Unknown Software")
    vendor = data.get("software_vendor", "Unknown Vendor")
    today = str(datetime.now().date())

    filename = f"{software.replace(' ', '_')}_security_review.pdf"
    pdf = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Title
    title_style = ParagraphStyle(
        name="TitleStyle",
        fontName="Helvetica-Bold",
        fontSize=20,
        alignment=TA_CENTER,
        spaceAfter=12
    )
    story.append(Paragraph("Software Security Review", title_style))

    # Centered subtext (software, vendor, date)
    sub_style = ParagraphStyle(
        name="SubStyle",
        fontName="Helvetica",
        fontSize=10,
        alignment=TA_CENTER,
        spaceAfter=12
    )
    story.append(Paragraph(f"Software: {software}", sub_style))
    story.append(Paragraph(f"Vendor: {vendor}", sub_style))
    story.append(Paragraph(f"Date: {today}", sub_style))

    # Calculate usable width for percentage-based columns (30/70)
    usable_width = pdf.width  # SimpleDocTemplate exposes usable width
    col_widths = [0.30 * usable_width, 0.70 * usable_width]

    body_style = styles["BodyText"]

    def section(title, keys):
        story.append(Paragraph(f"<b>{title}</b>", styles['Heading3']))
        story.append(Spacer(1, 4))

        # Header row as Paragraphs so they wrap too
        rows = [
            [Paragraph("<b>Question</b>", body_style), Paragraph("<b>Answer</b>", body_style)]
        ]

        for k in keys:
            raw_answer = data.get(k, "")
            # Normalize yes/no
            if k in YES_NO_QUESTIONS and isinstance(raw_answer, str):
                low = raw_answer.strip().lower()
                if low in ("y", "yes"):
                    raw_answer = "Y"
                elif low in ("n", "no"):
                    raw_answer = "N"
            # Ensure string and preserve line breaks with <br/>
            ans_text = "" if raw_answer is None else str(raw_answer)
            ans_text = ans_text.replace("\n", "<br/>")
            q_text = k.replace("_", " ").title()
            rows.append([
                Paragraph(q_text, body_style),
                Paragraph(ans_text, body_style)
            ])

        table = Table(rows, colWidths=col_widths, hAlign='LEFT')
        table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.lightblue),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(table)
        story.append(Spacer(1, 12))

    # Sections (with updated General Information)
    section("General Information", ["software_name", "software_vendor", "version", "use_case", "overview"])
    section("Policies & Data Handling",
            ["privacy_policy", "privacy_notes", "tos_reviewed", "tos_notes",
             "data_sharing", "data_sharing_notes", "data_collected", "data_storage"])
    section("Security Practices",
            ["pentesting", "pentest_details", "certifications", "auth_controls", "encryption"])
    section("Defender Scores",
            ["def_general", "def_security", "def_compliance", "def_notes"])
    section("Integrations",
            ["integrations", "integration_data", "integration_security"])
    section("Risk & Approval",
            ["risk_level", "key_risks", "mitigations", "approval", "approver", "approval_date"])

    pdf.build(story)
    print(f"PDF exported as {filename}")

# ------------------- WORD REPORT (table-style) -------------------

def generate_blank_word_report():
    """Generate a blank Word doc with section tables (Question | Answer)."""
    doc = Document()
    doc.add_heading("Software Security Review", 0)

    sections = {
        "General Information": ["software_name", "software_vendor", "version", "use_case", "overview"],
        "Policies & Data Handling": ["privacy_policy", "privacy_notes", "tos_reviewed", "tos_notes",
                                     "data_sharing", "data_sharing_notes", "data_collected", "data_storage"],
        "Security Practices": ["pentesting", "pentest_details", "certifications", "auth_controls", "encryption"],
        "Defender Scores": ["def_general", "def_security", "def_compliance", "def_notes"],
        "Integrations": ["integrations", "integration_data", "integration_security"],
        "Risk & Approval": ["risk_level", "key_risks", "mitigations", "approval", "approver", "approval_date"]
    }

    qdict = dict(QUESTIONS)

    for sec, keys in sections.items():
        doc.add_heading(sec, level=1)

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Question"
        hdr_cells[1].text = "Answer"
        for cell in hdr_cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(11)

        for k in keys:
            row_cells = table.add_row().cells
            row_cells[0].text = qdict.get(k, k)
            row_cells[1].text = ""
            for p in row_cells[0].paragraphs + row_cells[1].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

        doc.add_paragraph("\n")

    filename = "Blank_Software_Security_Review.docx"
    doc.save(filename)
    print(f"Blank Word report created: {filename}")

# ------------------- MAIN MENU -------------------



BLURB = "# SEC_REVIEW: CLI tool for quick software security reviews."

def main():
    try:
        while True:
            print("")
            print(BLURB)
            print("\n1. Start New Review")
            print("2. Continue Existing Review")
            print("3. Generate PDF Report")
            print("4. Generate Blank Word Report")
            print("5. Exit")

            choice = safe_input("Select an option: ")
            if choice is None:
                continue

            if choice == "1":
                start_review()
            elif choice == "2":
                continue_review()
            elif choice == "3":
                generate_pdf()
            elif choice == "4":
                generate_blank_word_report()
            elif choice == "5":
                print("")
                print("***Goodbye***")
                break
            else:
                print("*** Invalid selection, try again. ***")

    except KeyboardInterrupt:
        print("\033[1;31m\nOperation cancelled by user. Exiting gracefully.\033[0m")
        sys.exit(0)

if __name__ == "__main__":
    main()
